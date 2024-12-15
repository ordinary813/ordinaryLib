import os
import pandas as pd
from docx import Document
import re
import argparse
import json

# Map Hebrew numbers to integer values for words
hebrew_numbers = {
    "אחד": 1, "אחת": 1, "שתיים": 2, "שנים": 2, "שניים": 2, "שתים": 2, "שלוש": 3, "ארבע": 4, "חמש": 5, "שש": 6,
    "שבע": 7, "שמונה": 8, "תשע": 9, "עשר": 10, "עשרים": 20,
    "שלושים": 30, "ארבעים": 40, "חמישים": 50, "שישים": 60, "שבעים": 70, "שמונים": 80, "תשעים": 90,
    "מאה": 100, "מאתיים": 200, "אלף": 1000, "אלפיים": 2000
}


def hebrew_to_number(hebrew_text):
    # Remove leading "ה", "ו" and "-" characters if present
    parts = hebrew_text.split('-')
    for i in range(len(parts)):
        parts[i] = parts[i].lstrip("ה").lstrip("ו")

    total = 0
    i = 0

    while i < len(parts):
        part = parts[i]

        # Handle cases for number terms (like "חמש", "שלוש")
        if part in hebrew_numbers:
            # Handle "עשרה" as part of a number (e.g., "חמש עשרה" -> 15)
            if i + 1 < len(parts) and (parts[i + 1] == "עשרה" or parts[i + 1] == "עשר"):
                total += hebrew_numbers[part] + 10  # Add 10 to the number
                i += 1  # Skip over "עשרה"
            # Handle "מאות" (hundreds)
            elif i + 1 < len(parts) and parts[i + 1] == "מאות":
                total += hebrew_numbers[part] * 100  # Multiply by 100
                i += 1  # Skip over "מאות"
            else:
                total += hebrew_numbers[part]  # Add regular number
        i += 1  # Move to the next part

    return total


class Protocol:
    def __init__(self, file_name, file_path):
        # Initialize with details extracted from the file name
        self.file_path = file_path
        self.file_name = file_name
        self.kneset_number = self.extract_kneset_num(file_name)
        self.protocol_type = self.extract_protocol_type(file_name)
        self.protocol_text = self.load_document(file_path)
        self.protocol_number = self.extract_protocol_number()
        self.speaker_data = self.extract_speaker_data()

    def extract_kneset_num(self, file_name):
        underscore_index = file_name.find("_")
        if underscore_index != -1:
            try:
                kneset_number = int(file_name[:underscore_index])
                return kneset_number
            except Exception as e:
                print(f"Cannot find Knesset number in '{file_name}'")
        return -1

    def extract_protocol_type(self, file_name):
        underscore_index = file_name.rfind("_")  # last underscore
        if underscore_index != -1 and underscore_index > 0:
            try:
                protocol_type = file_name[underscore_index - 1:underscore_index]
                if protocol_type == "m":
                    return "plenary"
                if protocol_type == "v":
                    return "committee"
            except Exception as e:
                print(f"Error while extracting protocol type from '{file_name}'")
        return ""

    def load_document(self, file_path):
        try:
            doc = Document(file_path)
            return doc
        except Exception as e:
            print(f"Error loading document '{file_path}'")
        return ""

    def extract_protocol_number(self):
        protocol_number = -1

        # First check if it's a "committee" protocol
        if self.protocol_type == "committee":
            protocol_number = self.extract_committee_protocol_number()

        # If it's a "plenary" protocol, extract using a different term
        elif self.protocol_type == "plenary":
            protocol_number = self.extract_plenary_protocol_number()

        return protocol_number

    def extract_committee_protocol_number(self):
        # Look for the term "'פרוטוקול מס'" and the number after it
        protocol_number = -1
        for para in self.protocol_text.paragraphs:
            if "פרוטוקול מס" in para.text:
                # find the number after "פרוטוקול מס'"
                match = re.search(r"פרוטוקול מס'\s*(\d+)", para.text)
                if match:
                    protocol_number = int(match.group(1))  # convert the number to an integer
                    break
        return protocol_number

    def extract_plenary_protocol_number(self):
        # Look for the word "הכנסת" and extract the Hebrew number after it
        for para in self.protocol_text.paragraphs:
            if "הכנסת" in para.text:
                # Find the Hebrew number after "הכנסת"
                match = re.search(r"הישיבה\s+([^\s]+(?:-[^\s]+)*)\s+של הכנסת", para.text)
                if match:
                    hebrew_text = match.group(1)  # Extract the Hebrew number
                    return hebrew_to_number(hebrew_text)  # Convert Hebrew number to integer
        return -1

    def extract_speaker_data(self):
        document = self.protocol_text
        collected_speakers = []
        speaker_name, accumulated_text = None, ""

        for para in document.paragraphs:
            # Skip centered text and end of meeting
            if para.alignment == 1 or (para.style and para.style.name == 'Heading 1'):
                continue
            if "הישיבה ננעלה בשעה" in para.text:
                break

            # Remove brackets and marks from speaker name
            if re.search(r'<<[^>]+>>', para.text):
                para.text = re.sub(r'<<[^>]+>>', '', para.text).strip()

            if re.search(r'<[^>]+:>', para.text):
                para.text = re.sub(r'<(.*?)>', r'\1', para.text).strip()

            if self.protocol_type == 'plenary':  # ptm protocol
                if para.runs and (
                        (para.style and para.style.base_style and para.style.base_style.font.underline) or
                        (para.style and para.style.font.underline)
                ) and (
                        (para.style and para.style.base_style and para.style.base_style.font.bold) or
                        (para.style and para.style.font.bold)
                ) and (para.runs[-1].text == ':' or para.text.endswith(':')):

                    # Store the previous speaker and text if already exist
                    if speaker_name:
                        collected_speakers.append((speaker_name, accumulated_text.strip()))

                    # Process new speaker
                    speaker_name = para.text.replace(":", "").strip()
                    speaker_name = re.sub(r'\([^)]*\)', '', speaker_name).strip()
                    if len(speaker_name.split()) <= 1:
                        speaker_name = None
                        continue
                    speaker_name = " ".join(speaker_name.split()[-2:])  # Save only the last two words to avoid roles
                    accumulated_text = ""
                else:
                    # Accumulate text for the current speaker
                    if para.runs and not (
                            (para.style and para.style.base_style and para.style.base_style.font.underline) or
                            (para.style and para.style.font.underline)):
                        accumulated_text += para.text

            else:  # committee protocol type (ptv)
                if ((para.runs and (para.runs[-1].text == ':' or para.text.endswith(':'))) and
                    (para.runs[0].underline or (para.style and para.style.font and para.style.font.underline))
                ) or (para.style and para.style.base_style and para.style.base_style.font.underline):
                    # Store the previous speaker and text if already exist
                    if speaker_name:
                        collected_speakers.append((speaker_name, accumulated_text.strip()))

                    # Process new speaker
                    speaker_name = para.text.replace(":", "").strip()
                    speaker_name = re.sub(r'\([^)]*\)', '', speaker_name).strip()
                    if len(speaker_name.split()) <= 1:
                        speaker_name = None
                        continue
                    speaker_name = " ".join(speaker_name.split()[-2:])  # Save only the last two words to avoid roles
                    accumulated_text = ""
                else:
                    # Accumulate text for the current speaker
                    if para.runs and not (
                            (para.style and para.style.base_style and para.style.base_style.font.underline) or
                            (para.style and para.style.font.underline)):
                        accumulated_text += para.text

        # Add the last speaker and text to the list
        accumulated_text \
            = re.sub(r'<(.*?)>', '', accumulated_text).strip()
        if speaker_name:
            collected_speakers.append((speaker_name, accumulated_text.strip()))

        return collected_speakers


def clean_and_tokenize_sentences(text):
    # Pattern to match the end of a sentence (., !, ?)
    sentence_pattern = re.compile(r'(?<=[.!?])\s*')

    # Separate the text into sentences
    sentences = sentence_pattern.split(text.strip())

    # List to hold clean, tokenized sentences
    tokenized_sentences = []

    for sentence in sentences:
        sentence = sentence.strip()

        # Skip empty sentences
        if not sentence:
            continue

        # exclude sentences with specific unwanted symbols (– – – or – –)
        if "– – –" in sentence or "– –" in sentence:
            continue

        # exclude sentences with numbers
        if re.search(r'[a-zA-Z]', sentence):
            continue
        
        sentence = clean_sentence(sentence)
        # Tokenize the sentence
        tokens = tokenize_sentence(sentence)

        # Only include sentences with at least 4 tokens
        if len(tokens) >= 4:
            tokenized_sentences.append(tokens)

    return tokenized_sentences


def tokenize_sentence(sentence):
    pattern = r"""
        \d{1,2}[-/]\d{1,2}[-/]\d{2,4}        # Dates like 22/12/1992 or 12-05-21
        | \b(?:[A-Za-z]\.){2,}[A-Za-z]?      # English abbreviations like U.S.A., e.g.
        | \b[א-ת]+(?:'[\w"]+)?               # Hebrew abbreviations like ח"כ, היו"ר
        | \b[א-ת]'(?:\s?[א-ת]+)?            # Hebrew short names like ח' גולדשטיין
        | \b[א-ת]+\b                         # Hebrew words
        | \d+                                # Numbers
        | [.,!?;:\-\']                       # Punctuation
    """
    return re.findall(pattern, sentence, re.VERBOSE)

def clean_sentence(sentence):
    # Remove extra spaces before punctuation
    sentence = re.sub(r'\s+([.,!?])', r'\1', sentence)
    # Ensure the sentence ends with a word followed by a dot
    sentence = re.sub(r'\s+\.', '.', sentence.strip())
    return sentence




def produce_corpus(dir_path, write_to_txt=False):
    df = pd.DataFrame(columns=["protocol_name",
                               "knesset_number",
                               "protocol_type",
                               "protocol_number",
                               "speaker_name",
                               "sentence_text"
                               ])
    docx_count = 0
    docx_processed = 0

    for _, _, files in os.walk(dir_path):
        for file in files:
            if file.endswith('.docx'):
                docx_count += 1

    print(f"{docx_processed}/{docx_count} docx files processed.", end='\r')

    for curr_file in os.listdir(dir_path):
        if curr_file.endswith('.docx'):
            file_path = os.path.join(dir_path, os.fsdecode(curr_file))
            protocol = Protocol(curr_file, file_path)
            data = protocol.speaker_data

            # Convert tuples to dictionaries
            data = [{"speaker_name": speaker_name, "sentence_text": accumulated_text} for speaker_name, accumulated_text
                    in data]

            # Add additional protocol information to each entry
            for entry in data:
                entry["protocol_name"] = curr_file
                entry["knesset_number"] = protocol.kneset_number
                entry["protocol_type"] = protocol.protocol_type
                entry["protocol_number"] = protocol.protocol_number

            temp_df = pd.DataFrame(data)

            # sentences tokenization
            expanded_data = []
            for _, row in temp_df.iterrows():
                speech = row["sentence_text"]
                tokenized_sentences = clean_and_tokenize_sentences(speech)

                # For each tokenized sentence, create a new entry
                for tokens in tokenized_sentences:
                    new_entry = row.to_dict()
                    new_entry["sentence_text"] = " ".join(tokens)  # Convert tokens back into sentence string
                    expanded_data.append(new_entry)

            expanded_df = pd.DataFrame(expanded_data)
            df = pd.concat([df, expanded_df], ignore_index=True)

            print(f"{docx_processed}/{docx_count} docx files processed.", end='\r', flush=True)
            docx_processed += 1

    print(f"All docx files processed.", flush=True)
    return df


def save_dataframe_to_jsonl(df, output_path):
    with open(output_path, 'w', encoding='utf-8') as jsonl_file:
        for _, row in df.iterrows():
            json_object = row.to_dict()
            jsonl_file.write(f"{json.dumps(json_object, ensure_ascii=False)}\n")


# to run the script in the command line, use the following syntax:
# 'python processing_knesset_corpus.py <path/to/protocols> <path/to/output/file>'
def main():
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "directory",
        type=str,
        help="Path to the directory containing protocol files."
    )

    parser.add_argument(
        "output_path",
        type=str,
        help="Path to the directory for the output JSONL."
    )

    args = parser.parse_args()

    dir_path = args.directory
    out_path = args.output_path

    print("Processing Protocols...")
    df = produce_corpus(dir_path, write_to_txt=True)

    save_dataframe_to_jsonl(df, out_path)
    print(f"Corpus was saved to {out_path}.")


if __name__ == "__main__":
    main()
